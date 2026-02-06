import logging
from pathlib import Path
from typing import Dict, Tuple, Optional

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import Contrato
from .utils import formatar_valor, valor_por_extenso, formatar_data
from .config import get_config

logger = logging.getLogger(__name__)


class DocumentGenerator:
    def __init__(self, template_path: str, prints_dir: str = None):
        self.template_path = Path(template_path)
        self.prints_dir = Path(prints_dir) if prints_dir else None
        self._template_cache = None

    def _carregar_template(self) -> Document:
        return Document(self.template_path)

    def _buscar_imagem_contrato(self, numero_contrato: str) -> Optional[Path]:
        if not self.prints_dir or not self.prints_dir.exists():
            return None

        for ext in [".png", ".jpg", ".jpeg", ".PNG", ".JPG", ".JPEG"]:
            img_path = self.prints_dir / f"{numero_contrato}{ext}"
            if img_path.exists():
                return img_path

        return None

    def _inserir_imagem_no_paragrafo(
        self, para, img_path: Path, width_inches: float = 5.5
    ) -> bool:

        try:
            for run in para.runs:
                run.text = ""
            run = para.add_run()
            run.add_picture(str(img_path), width=Inches(width_inches))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True

        except Exception as e:
            logger.warning(f"Erro ao inserir imagem: {e}")
            return False

    def _substituir_texto(self, doc: Document, substituicoes: Dict[str, str]) -> int:
        total = 0
        for para in doc.paragraphs:
            total += self._substituir_em_paragrafo(para, substituicoes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        total += self._substituir_em_paragrafo(para, substituicoes)
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    total += self._substituir_em_paragrafo(para, substituicoes)
            if section.footer:
                for para in section.footer.paragraphs:
                    total += self._substituir_em_paragrafo(para, substituicoes)

        return total

    def _substituir_em_paragrafo(self, para, substituicoes: Dict[str, str]) -> int:
        total = 0
        texto_completo = para.text
        texto_modificado = texto_completo

        for antigo, novo in substituicoes.items():
            if antigo in texto_modificado:
                texto_modificado = texto_modificado.replace(antigo, novo, 1)
                total += 1

        if texto_modificado != texto_completo:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = texto_modificado
            else:
                para.add_run(texto_modificado)

        return total

    def gerar(self, contrato: Contrato, output_path: str) -> Tuple[bool, str]:
        try:
            doc = self._carregar_template()
            subs = self._montar_substituicoes(contrato)
            total = self._substituir_texto(doc, subs)
            if contrato.valor_atualizado > 0:
                valor_atual_formatado = formatar_valor(contrato.valor_atualizado)
                valor_atual_extenso = valor_por_extenso(contrato.valor_atualizado)
                valor_atual_completo = (
                    f"{valor_atual_formatado} ({valor_atual_extenso})"
                )

                placeholder_valor = "R$XXXXXX (escrever o valor por extenso)"

                for para in doc.paragraphs:
                    if placeholder_valor in para.text:
                        para.text = para.text.replace(
                            placeholder_valor, valor_atual_completo, 1
                        )
                        total += 1
                        break

            img_path = self._buscar_imagem_contrato(contrato.numero)
            img_inserida = False
            if img_path:
                placeholder_clausula = "(transcrever ou printar a cláusula do contrato relativa ao pagamento)"
                for para in doc.paragraphs:
                    if placeholder_clausula in para.text:
                        img_inserida = self._inserir_imagem_no_paragrafo(para, img_path)
                        break

            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            msg = f"Gerado com {total} substituições"
            if img_inserida:
                msg += " + imagem"
            return True, msg

        except Exception as e:

            return False, str(e)

    def _montar_substituicoes(self, c: Contrato) -> Dict[str, str]:
        subs = {}
        config = get_config()
        if c.proprietarios:
            blocos_props = []
            for i, prop in enumerate(c.proprietarios):
                nacionalidade = "brasileiro(a)"
                bloco = f"{prop.nome.upper()}, {nacionalidade}, inscrito(a) no CPF sob o nº {prop.cpf or '(inserir o CPF do locador)'}"
                if prop.rg:
                    bloco += f" e no RG nº {prop.rg}"
                endereco = prop.endereco or "(incluir endereço completo do locador)"
                bloco += f", residente e domiciliado(a) à {endereco}"
                email = prop.email or "(inserir o e-mail)"
                bloco += f", com endereço eletrônico {email}"
                blocos_props.append(bloco)

            if len(blocos_props) > 1:
                bloco_completo_props = ", e ".join(blocos_props)
            else:
                bloco_completo_props = blocos_props[0]
            bloco_template_locador = ""
            subs[bloco_template_locador] = bloco_completo_props

        if c.inquilinos:
            inq1 = c.inquilinos[0]
            bloco_inq1 = (
                f"{inq1.nome.upper()}, ({inq1.nacionalidade}),  "
                f"inscrito(a) no CPF sob o n.º {inq1.cpf or '(inserir o CPF do Inquilino)'}, "
                f"Telefone {inq1.telefone or '(DDD) (número do whatsapp do Inquilino)'}, "
                f"e-mail(s) {inq1.email or '(inserir o endereço eletrônico do Inquilino)'}"
            )

            if len(c.inquilinos) > 1:
                inq2 = c.inquilinos[1]
                bloco_inq2 = (
                    f" e {inq2.nome.upper()}, ({inq2.nacionalidade}),  "
                    f"inscrito(a) no CPF sob o n.º {inq2.cpf or '(inserir o CPF do Inquilino)'}, "
                    f"Telefone {inq2.telefone or '(DDD) (número do whatsapp do Inquilino)'}, "
                    f"e-mail(s) {inq2.email or '(inserir o endereço eletrônico do Inquilino)'}"
                )
                bloco_completo = bloco_inq1 + bloco_inq2

            else:
                bloco_completo = bloco_inq1
            bloco_template = ""
            subs[bloco_template] = bloco_completo
        endereco_imovel = (
            c.imovel.endereco_completo
            if c.imovel.endereco_completo
            else "(inserir endereço do imóvel)"
        )

        subs[
            "(inserir o endereço completo do imóvel locado objeto do contrato: Rua/Avenida, número, complemento, Cidade, UF e CEP)"
        ] = endereco_imovel

        subs["(inserir o endereço completo dos Inquilinos)"] = endereco_imovel

        if c.valor_historico > 0:
            valor_hist_formatado = formatar_valor(c.valor_historico)
            valor_hist_extenso = valor_por_extenso(c.valor_historico)
            subs["R$XXXXXX (escrever o valor por extenso)"] = (
                f"{valor_hist_formatado} ({valor_hist_extenso})"
            )

        if c.valor_causa > 0:
            valor_formatado = formatar_valor(c.valor_causa)
            valor_extenso = valor_por_extenso(c.valor_causa)
            subs["R$00.000,00 (inserir o valor por extenso)"] = (
                f"{valor_formatado} ({valor_extenso})"
            )

        subs["Cidade, dia de mês de 2025."] = f"{c.cidade}, {formatar_data()}."
        subs["(DDD) XXXX-YYYY"] = config["escritorio_telefone"]
        subs["(DDD) 9XXXX-YYYY"] = config["escritorio_whatsapp"]
        subs["inserir o e-mail do escritório ou assessoria de cobrança"] = config[
            "escritorio_email"
        ]
        subs["(inserir o nome do advogado responsável do escritório)"] = config[
            "advogado_nome"
        ]
        subs["XXX.XXX"] = config["advogado_oab"]
        subs["(inserir o endereço comercial do escritório)"] = config[
            "escritorio_endereco"
        ]
        subs[
            "(inserir o e-mail oficial do escritório para recebimento de intimações)"
        ] = config.get("escritorio_email_intimacoes", config["escritorio_email"])
        return subs
