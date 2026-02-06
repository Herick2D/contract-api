
import re
import json
import uuid

from pathlib import Path
from datetime import datetime
from typing import List, Optional, Dict, Any
from docx import Document

from ..config import get_settings


class TemplateService:
    def __init__(self):
        self.settings = get_settings()
        self.templates_dir = self.settings.templates_dir
        self.metadata_file = self.templates_dir / "templates_metadata.json"
        self._ensure_metadata_file()

    def _ensure_metadata_file(self):
        if not self.metadata_file.exists():
            self._save_metadata({})

    def _load_metadata(self) -> Dict[str, Any]:
        try:
            with open(self.metadata_file, "r", encoding="utf-8") as f:
                return json.load(f)

        except:
            return {}

    def _save_metadata(self, metadata: Dict[str, Any]):
        with open(self.metadata_file, "w", encoding="utf-8") as f:
            json.dump(metadata, f, indent=2, ensure_ascii=False, default=str)

    def _extract_placeholders(self, file_path: Path) -> List[str]:
        placeholders = set()

        try:
            doc = Document(file_path)
            def find_placeholders(text: str):
                patterns = [
                    r"\([^)]{5,}\)",
                    r"XXXXX+",
                    r"\{[^}]+\}",
                ]

                found = []
                for pattern in patterns:
                    matches = re.findall(pattern, text)
                    found.extend(matches)
                return found

            for para in doc.paragraphs:
                if para.text.strip():
                    phs = find_placeholders(para.text)
                    placeholders.update(phs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                phs = find_placeholders(para.text)
                                placeholders.update(phs)
            for section in doc.sections:

                if section.header:
                    for para in section.header.paragraphs:
                        phs = find_placeholders(para.text)
                        placeholders.update(phs)
                if section.footer:
                    for para in section.footer.paragraphs:
                        phs = find_placeholders(para.text)
                        placeholders.update(phs)

        except Exception as e:

            print(f"Erro ao extrair placeholders: {e}")

        return list(placeholders)[:50]

    async def create_template(
        self,
        name: str,
        file_content: bytes,
        original_filename: str,
        description: Optional[str] = None,
    ) -> Dict[str, Any]:

        template_id = str(uuid.uuid4())[:8]
        file_ext = Path(original_filename).suffix
        new_filename = f"{template_id}_{name.replace(' ', '_')}{file_ext}"
        file_path = self.templates_dir / new_filename
        with open(file_path, "wb") as f:
            f.write(file_content)
        placeholders = self._extract_placeholders(file_path)

        template_data = {
            "id": template_id,
            "name": name,
            "description": description or "",
            "filename": original_filename,
            "file_path": str(file_path),
            "status": "active",
            "placeholders": placeholders,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(),
        }

        metadata = self._load_metadata()
        metadata[template_id] = template_data
        self._save_metadata(metadata)
        return template_data

    async def get_template(self, template_id: str) -> Optional[Dict[str, Any]]:
        metadata = self._load_metadata()
        return metadata.get(template_id)

    async def list_templates(
        self, status: Optional[str] = None
    ) -> List[Dict[str, Any]]:

        metadata = self._load_metadata()
        templates = list(metadata.values())
        if status:
            templates = [t for t in templates if t.get("status") == status]

        return sorted(templates, key=lambda x: x.get("created_at", ""), reverse=True)

    async def update_template(
        self,
        template_id: str,
        name: Optional[str] = None,
        description: Optional[str] = None,
        status: Optional[str] = None,
    ) -> Optional[Dict[str, Any]]:

        metadata = self._load_metadata()

        if template_id not in metadata:
            return None

        template = metadata[template_id]
        if name is not None:
            template["name"] = name

        if description is not None:
            template["description"] = description

        if status is not None:
            template["status"] = status

        template["updated_at"] = datetime.now().isoformat()
        metadata[template_id] = template
        self._save_metadata(metadata)

        return template

    async def delete_template(self, template_id: str) -> bool:
        metadata = self._load_metadata()
        if template_id not in metadata:
            return False

        template = metadata[template_id]
        file_path = Path(template["file_path"])
        if file_path.exists():
            file_path.unlink()
        del metadata[template_id]
        self._save_metadata(metadata)
        return True

    def get_template_path(self, template_id: str) -> Optional[Path]:
        metadata = self._load_metadata()
        if template_id not in metadata:
            return None

        file_path = Path(metadata[template_id]["file_path"])
        if not file_path.exists():
            return None

        return file_path

_template_service: Optional[TemplateService] = None


def get_template_service() -> TemplateService:
    global _template_service
    if _template_service is None:
        _template_service = TemplateService()
    return _template_service
